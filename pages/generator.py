"""AI 定制简历 & 求职信生成页面"""
import streamlit as st
import anthropic, json, io
from docx import Document
from docx.shared import Pt
from utils.database import get_jobs, get_base_resume, save_resume, save_application
from utils.i18n import t

RESUME_PROMPT = """You are a professional resume consultant.
Tailor the candidate's resume for the specific job description below.

Rules:
1. Highlight skills and experience most relevant to this JD
2. Incorporate JD keywords naturally without fabricating experience
3. Quantify achievements where possible
4. Keep all real information — optimize, don't invent
5. Output the complete tailored resume text

Resume:
{resume}

Target Job:
Title: {title}
Company: {company}
JD:
{jd}

Output the tailored resume directly, no preamble."""

COVER_PROMPT = """You are a professional career coach. Write a compelling cover letter.

Requirements:
- Max 350 words
- Open with the position and your top strength
- 2-3 paragraphs covering most relevant experience
- Close with enthusiasm and call to action
- Professional but not stiff
- Write in {language}

Resume summary:
{resume_summary}

Target role:
Title: {title}
Company: {company}
JD highlights: {jd_summary}

Output the cover letter body directly, no salutation needed."""

def gen_resume(api_key, resume, job):
    client = anthropic.Anthropic(api_key=api_key)
    with client.messages.stream(
        model="claude-sonnet-4-20250514", max_tokens=2000,
        messages=[{"role":"user","content":RESUME_PROMPT.format(
            resume=resume, title=job["title"], company=job["company"], jd=job["description"][:3000]
        )}]
    ) as s: return s.get_final_text()

def gen_cover(api_key, resume, job, language):
    client = anthropic.Anthropic(api_key=api_key)
    with client.messages.stream(
        model="claude-sonnet-4-20250514", max_tokens=800,
        messages=[{"role":"user","content":COVER_PROMPT.format(
            resume_summary=resume[:800], title=job["title"], company=job["company"],
            jd_summary=job["description"][:1000], language=language
        )}]
    ) as s: return s.get_final_text()

def make_docx(title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.split("\n"):
        p = doc.add_paragraph(line); p.style.font.size = Pt(11)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

def show():
    st.markdown(f"## {t('generator_title')}")
    api_key = st.session_state.get("api_key","")
    if not api_key: st.warning(t("no_api_key")); return
    base = get_base_resume()
    if not base: st.warning(t("no_base_resume")); return

    jobs = get_jobs(limit=100)
    scored = sorted([j for j in jobs if j["ai_score"] is not None], key=lambda x: x["ai_score"], reverse=True)
    all_options = scored + [j for j in jobs if j["ai_score"] is None]

    if not all_options: st.info(t("no_jobs_db")); return

    job_options = {
        f"{'🟢' if (j['ai_score'] or 0)>=70 else '🟡' if (j['ai_score'] or 0)>=50 else '⚪'} "
        f"{j['title']} @ {j['company']} ({j['source']})"
        f"{' — '+str(int(j['ai_score'])) if j['ai_score'] else ''}": j
        for j in all_options
    }

    default_idx = 0
    if "selected_job_id" in st.session_state:
        for i,j in enumerate(all_options):
            if j["id"] == st.session_state["selected_job_id"]:
                default_idx = i; break

    selected_label = st.selectbox(t("select_job"), list(job_options.keys()), index=default_idx)
    job = job_options.get(selected_label)
    if not job: st.warning("Please select a job"); return

    col1, col2 = st.columns(2)
    do_resume = col1.checkbox(t("gen_resume_cb"), value=True)
    do_cover  = col2.checkbox(t("gen_cover_cb"),  value=True)
    lang_opts = ["Chinese (中文)", "English"] if st.session_state.get("lang")=="en" else ["中文", "English"]
    cl_lang   = col2.selectbox(t("cover_lang"), lang_opts) if do_cover else lang_opts[0]

    if st.button(t("start_gen"), type="primary"):
        tailored = cover = ""
        if do_resume:
            with st.spinner("🤖 Generating tailored resume..."):
                try: tailored = gen_resume(api_key, base["content"], job); st.success("✅ Done")
                except Exception as e: st.error(str(e))
        if do_cover:
            with st.spinner("✉️ Writing cover letter..."):
                try: cover = gen_cover(api_key, base["content"], job, cl_lang)
                except Exception as e: st.error(str(e))

        if tailored:
            st.markdown(t("tailored_resume"))
            edited = st.text_area(t("editable"), value=tailored, height=500)
            c1, c2 = st.columns(2)
            if c1.button(t("save_version")):
                name = f"{job['company']} - {job['title']}"
                rid  = save_resume(name, edited, is_base=False)
                st.session_state["last_resume_id"] = rid
                st.success(t("saved_as", name))
            c2.download_button(t("download_word"), make_docx(job["title"], edited),
                file_name=f"resume_{job['company']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        if cover:
            st.markdown(t("cover_letter"))
            edited_cl = st.text_area(t("editable"), value=cover, height=300, key="cl_edit")
            st.download_button(t("download_cover"), edited_cl.encode("utf-8"),
                file_name=f"cover_{job['company']}.txt", mime="text/plain")

        if tailored or cover:
            st.markdown("---")
            if st.button(t("mark_applied"), type="primary"):
                rid = st.session_state.get("last_resume_id")
                if not rid:
                    rid = save_resume(f"{job['company']}_{job['title']}", tailored or base["content"], is_base=False)
                save_application(job["id"], rid, cover)
                st.success(t("applied_success", job["title"], job["company"]))
